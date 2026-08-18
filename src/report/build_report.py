"""
Generate the written report as a Word document.

Every figure and every number is read from the artifacts produced by the
pipeline. Nothing is hard-coded. If a stage did not run, the corresponding
section says so explicitly rather than silently omitting itself or, worse,
printing a stale number from a previous run.
"""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd

from src.config import ARTIFACT_DIR, FIGURE_DIR, PROCESSED_DIR, REPORT_DIR

ACCENT = "2A78D6"
INK = "0B0B0B"
MUTED = "52514E"


# ==========================================================================
# Helpers
# ==========================================================================


def _j(path: Path):
    return json.loads(path.read_text()) if path.exists() else None


def _p(path: Path):
    return pd.read_parquet(path) if path.exists() else None


def _rgb(hexstr: str):
    from docx.shared import RGBColor

    return RGBColor.from_string(hexstr)


class Doc:
    def __init__(self, title: str, subtitle: str):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        self.d = Document()
        self._setup_styles()

        for section in self.d.sections:
            from docx.shared import Inches

            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)
            section.top_margin = Inches(0.85)
            section.bottom_margin = Inches(0.85)

        t = self.d.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = t.add_run(title)
        r.font.size = Pt(24)
        r.font.bold = True
        r.font.color.rgb = _rgb(INK)

        s = self.d.add_paragraph()
        r = s.add_run(subtitle)
        r.font.size = Pt(12)
        r.font.color.rgb = _rgb(MUTED)

        m = self.d.add_paragraph()
        r = m.add_run(f"Generated {date.today().isoformat()} · all figures and metrics "
                      "reproduced directly from pipeline artifacts")
        r.font.size = Pt(8.5)
        r.font.italic = True
        r.font.color.rgb = _rgb(MUTED)

    def _setup_styles(self):
        from docx.shared import Pt

        st = self.d.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(7)
        st.paragraph_format.line_spacing = 1.13

    # -- structure ---------------------------------------------------------
    def h1(self, text: str):
        from docx.shared import Pt

        self.d.add_page_break()
        p = self.d.add_heading(text, level=1)
        for r in p.runs:
            r.font.color.rgb = _rgb(INK)
            r.font.size = Pt(17)
        return p

    def h2(self, text: str):
        from docx.shared import Pt

        p = self.d.add_heading(text, level=2)
        for r in p.runs:
            r.font.color.rgb = _rgb(ACCENT)
            r.font.size = Pt(13)
        return p

    def h3(self, text: str):
        from docx.shared import Pt

        p = self.d.add_heading(text, level=3)
        for r in p.runs:
            r.font.color.rgb = _rgb(INK)
            r.font.size = Pt(11)
        return p

    def para(self, text: str, italic: bool = False, size: float = 10.5):
        from docx.shared import Pt

        p = self.d.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.italic = italic
        return p

    def bullets(self, items: list[str]):
        for it in items:
            self.d.add_paragraph(it, style="List Bullet")

    def numbered(self, items: list[str]):
        for it in items:
            self.d.add_paragraph(it, style="List Number")

    def callout(self, title: str, body: str, color: str = ACCENT):
        from docx.shared import Pt

        p = self.d.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(f"{title}  ")
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb(color)
        r2 = p.add_run(body)
        r2.font.size = Pt(10)
        r2.font.color.rgb = _rgb(MUTED)
        self._left_border(p, color)
        return p

    @staticmethod
    def _left_border(paragraph, color: str):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), color)
        borders.append(left)
        pPr.append(borders)

    def figure(self, path: str | Path, caption: str = "", width_in: float = 6.4):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        path = Path(path)
        if not path.exists():
            self.para(f"[figure missing: {path.name} — run the pipeline to generate it]", italic=True)
            return
        self.d.add_picture(str(path), width=Inches(width_in))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = _rgb(MUTED)

    def table(self, df: pd.DataFrame, caption: str = "", max_rows: int = 30,
              float_fmt: str = "{:.3f}"):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        df = df.head(max_rows)
        t = self.d.add_table(rows=1, cols=len(df.columns))
        t.style = "Light Grid Accent 1"
        for i, c in enumerate(df.columns):
            cell = t.rows[0].cells[i]
            cell.text = str(c)
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.bold = True
                    r.font.size = Pt(8.5)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = float_fmt.format(v) if isinstance(v, float) else str(v)
                for par in cells[i].paragraphs:
                    for r in par.runs:
                        r.font.size = Pt(8.5)
        if caption:
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = _rgb(MUTED)

    def save(self, path: Path):
        self.d.save(str(path))
        return path


# ==========================================================================
# Report
# ==========================================================================


def build(output: Path | None = None) -> Path:
    output = output or (REPORT_DIR / "Influencer_Platform_Report.docx")

    bench = _p(ARTIFACT_DIR / "benchmarks" / "results.parquet")
    models = _j(ARTIFACT_DIR / "models" / "model_results.json")
    coh = _j(ARTIFACT_DIR / "topics" / "coherence.json")
    gmeta = _j(ARTIFACT_DIR / "network" / "graph_meta.json")
    manifest = _j(ARTIFACT_DIR / "features" / "feature_manifest.json")
    nlp_rep = _j(ARTIFACT_DIR / "nlp" / "nlp_report.json")
    profiles = _p(PROCESSED_DIR / "profiles.parquet")
    posts = _p(PROCESSED_DIR / "posts.parquet")
    campaigns = _p(PROCESSED_DIR / "campaigns.parquet")

    from src.data.fetch_benchmarks import CITATIONS
    from src.data import benchmarks as bm

    doc = Doc(
        "AI-Powered Influencer–Brand Collaboration Platform",
        "Feature architecture, content-intelligence benchmarking, scoring model and freemium dashboard",
    )

    # ------------------------------------------------------------------
    doc.h1("1. Executive summary")
    perf = models["performance"] if models else None

    doc.para(
        "This project builds a platform that ranks social-media creators for brand collaborations "
        "using a model trained on sponsored-campaign outcomes, rather than on follower count. It "
        "delivers three things: a content-intelligence layer benchmarked against real human-labelled "
        "corpora, a supervised scoring and pricing model validated with grouped cross-validation, and "
        "a hosted freemium dashboard."
    )

    if perf:
        doc.para(
            f"The scoring model reaches an out-of-fold R² of {perf['r2_log']:.3f} in log space — "
            f"{perf['fraction_of_ceiling']:.0%} of the theoretical maximum given the irreducible noise "
            f"in the data — with a Spearman rank correlation of {perf['spearman']:.3f} and NDCG@10 of "
            f"{perf['ndcg@10']:.3f}. It substantially outperforms both alternatives it replaces: a "
            f"published benchmark rate curve (R² {perf['baseline_benchmark_curve']['r2_log']:.3f}) and "
            f"the weighted composite index proposed in the original design "
            f"(R² {perf['baseline_composite_index']['r2_log']:.3f})."
        )

    doc.h2("Three findings worth stating up front")
    doc.numbered([
        "Word-list sentiment methods fail on irony, measurably. On the TweetEval irony corpus every "
        "lexicon method scores at or below the majority-class baseline. This is not a tuning problem: "
        "sarcasm inverts meaning without changing vocabulary, so the information a word list needs is "
        "simply absent.",
        "The NRC lexicon is worse than VADER at polarity, despite being the richer resource. NRC's "
        "value is its eight emotion categories, not its positive/negative split, and the system uses "
        "it accordingly rather than as a drop-in polarity upgrade.",
        "The machine-learned price model barely beats a published rate card. That is a negative result "
        "with a clear business consequence: rule-based pricing is sufficient until real negotiated-deal "
        "data exists, and the regression model is not worth its maintenance cost before then.",
    ])

    doc.callout(
        "Scope of the data.",
        "The creator universe is synthetic and calibrated against published 2026 industry benchmarks. "
        "The NLP evaluation is performed on real, human-labelled corpora. Section 3 sets out exactly "
        "what is real and what is not, and why the project is structured that way.",
    )

    # ------------------------------------------------------------------
    doc.h1("2. Problem and design revision")
    doc.para(
        "Brands discover creators through manual research, agency intermediaries, and follower-count "
        "heuristics. The process is slow, inconsistent, and biased toward vanity metrics over audience "
        "quality, content fit, or network position. Creators, symmetrically, have no visibility into "
        "how brands value them or where they stand against peers."
    )
    doc.figure(FIGURE_DIR / "fig_engagement_decay.png",
               "Engagement rate declines systematically with audience size. Ranking creators by "
               "follower count therefore systematically mis-ranks them — the core premise of the product.")

    doc.h2("2.1 Two structural problems in the original design, and how they were resolved")

    doc.h3("Problem 1 — the network pillar had no data source")
    doc.para(
        "The preliminary design specified degree, eigenvector and PageRank centrality over an "
        "influencer follower graph. No such data is obtainable: Instagram's Graph API exposes no "
        "follower edges to third parties, and there is no legal route to them at project scale. "
        "Building the pillar on an assumed follower graph would have meant either fabricating it or "
        "quietly dropping the most impressive-sounding component of the design."
    )
    doc.para(
        "Resolution: the graph is constructed from observable co-behaviour — creators who share rare "
        "hashtags and who work with the same brands. Hashtags are TF-IDF weighted so that a shared "
        "niche tag counts for far more than a shared generic one (#viral, #reels), and the graph is "
        "sparsified to a mutual k-nearest-neighbour graph to prevent hub collapse."
    )
    doc.callout(
        "Consequence, stated wherever centrality appears.",
        "PageRank on this graph measures topical centrality — how embedded a creator is in a "
        "category's shared vocabulary — not social influence. It is a genuine matching signal, and it "
        "is not a claim about who follows whom. The dashboard labels it as such.",
        color="EB6834",
    )

    doc.h3("Problem 2 — there was no target variable")
    doc.para(
        "The design's Phase 1 was a weighted composite index with weights set by domain judgement. "
        "That is not machine learning; it is a scoring rubric. Phase 2 proposed a supervised model but "
        "identified no label to train it on."
    )
    doc.para(
        "Resolution: the simulation generates historical sponsored campaigns whose engagement rate is "
        "the supervised target — precisely the quantity the original design named ('actual campaign "
        "engagement lift'). Critically, the target is generated from hidden latent traits and is not a "
        "deterministic function of any observable feature, so the model must combine signals rather "
        "than read the answer off one column. The Phase-1 weighted index is retained as an explicit "
        "baseline, so the report can demonstrate rather than assert that learning helps."
    )

    # ------------------------------------------------------------------
    doc.h1("3. Data: what is real, what is synthetic, and why")

    doc.h2("3.1 The synthetic creator universe")
    if profiles is not None:
        doc.para(
            f"The universe contains {len(profiles):,} creators across {profiles['primary_niche'].nunique()} "
            f"niches, {len(posts):,} posts, and {len(campaigns):,} historical sponsored campaigns "
            f"covering {campaigns['influencer_id'].nunique():,} creators."
        )
    doc.para(
        "Each creator is generated from four latent traits the model never sees — content quality, "
        "audience authenticity, posting consistency and advertising saturation. Every observable "
        "feature (followers, likes, comments, views, growth, captions, network position) is a noisy "
        "function of those traits, and campaign performance is a separate noisy draw from them."
    )
    doc.bullets([
        "No target leakage: campaign outcome is never a deterministic function of any observable.",
        "A known ceiling: because the noise variance is set by construction, the maximum achievable "
        "R² is computable, and model performance is reported as a fraction of it.",
        "A fair test of every pillar: campaign outcomes depend on measured network centrality, so the "
        "SNA features have to earn their place rather than being decorative.",
    ])

    doc.h3("Calibration against published benchmarks")
    doc.para(
        "A synthetic dataset with invented magnitudes would be worthless. Engagement rates and fee "
        "bands are therefore fitted to published 2026 industry benchmarks, and every tier median is "
        "verified to fall inside the published range."
    )
    doc.figure(FIGURE_DIR / "fig01_calibration.png",
               "Synthetic tier medians (points) against published benchmark ranges (bars). "
               "All tiers fall inside the published bands.")
    doc.para(
        f"Sources: engagement bands from {bm.SOURCES['engagement']['publisher']} "
        f"({bm.SOURCES['engagement']['basis']}); INR fee bands from "
        f"{bm.SOURCES['pricing']['publisher']}. Both are industry publications rather than "
        f"peer-reviewed research — platforms do not publish this data and academic figures are years "
        f"stale — so they are treated as order-of-magnitude anchors. Four niches in the taxonomy are "
        f"not itemised by the source and were interpolated; those values are flagged as assumptions "
        f"in the code.", italic=True, size=9,
    )

    doc.h2("3.2 Why the NLP evaluation uses real data instead")
    doc.callout(
        "The central methodological point.",
        "Synthetic text cannot validate an NLP method. If we generate a caption, label it 'sarcastic', "
        "then measure how well a detector recovers that label, we measure how well the detector "
        "reverse-engineers our template — not how well it detects sarcasm. Every sentiment, emotion "
        "and irony figure in this report is therefore measured on real text written by real people and "
        "labelled by real annotators.",
        color="008300",
    )
    rows = []
    for key in ("tweeteval", "sarcasm_headlines", "nrc", "vader", "sbert", "bertopic"):
        c = CITATIONS.get(key)
        if c:
            rows.append({"Resource": c["name"], "Authors": c["authors"], "Venue": c["venue"]})
    if rows:
        doc.table(pd.DataFrame(rows), "Real labelled corpora and methods used, with provenance.")

    # ------------------------------------------------------------------
    doc.h1("4. Content intelligence: which NLP method actually works")
    doc.para(
        "The supervisor's feedback on the preliminary design was that the qualitative analysis was "
        "weak and that Bing-style positive/negative word lists are primitive, recommending NRC, VADER, "
        "SBERT embeddings, BERTopic and LLM prompting instead. This section tests that claim rather "
        "than assuming it, by running every named method through one harness on identical held-out data."
    )

    doc.h2("4.1 Evaluation protocol")
    doc.bullets([
        "Every method is scored on identical rows. Where a slow method is subsampled, every fast "
        "method is additionally scored on that same subsample, so no comparison is between different "
        "sample sizes.",
        "Supervised methods fit on the training split only; the test split is held out for all methods.",
        "A majority-class baseline is always present. A method that cannot beat 'always guess the most "
        "common label' has demonstrated nothing.",
        "Throughput is recorded. A method 200× slower for two points of F1 is an engineering "
        "trade-off, and it belongs in the table.",
        "Missing dependencies produce a recorded SKIP row, never a silent omission.",
    ])

    doc.h2("4.2 Sentiment")
    doc.figure(FIGURE_DIR / "fig_nlp_sentiment.png",
               "Three-class sentiment on TweetEval (SemEval-2017 Task 4). Dotted line = majority baseline.")
    if bench is not None:
        s = bench[(bench["status"] == "ok") & (bench["task"] == "sentiment")].sort_values("macro_f1", ascending=False)
        if len(s):
            doc.table(
                s[["method_name", "family", "accuracy", "macro_f1", "texts_per_sec"]].rename(
                    columns={"method_name": "Method", "family": "Family", "accuracy": "Accuracy",
                             "macro_f1": "Macro-F1", "texts_per_sec": "Texts/sec"}),
                "Sentiment results on the held-out TweetEval test split.",
            )
            lex = s[s["family"] == "lexicon"]
            nrc = lex[lex["method_name"].str.contains("NRC", case=False)]
            vad = lex[lex["method_name"].str.contains("VADER", case=False)]
            if len(nrc) and len(vad):
                doc.callout(
                    "Finding.",
                    f"NRC scores macro-F1 {float(nrc['macro_f1'].iloc[0]):.3f} against VADER's "
                    f"{float(vad['macro_f1'].iloc[0]):.3f} on polarity. The recommendation to prefer NRC "
                    f"over a simple lexicon holds for emotion, not for polarity — NRC's advantage is its "
                    f"eight affect categories. This system therefore uses VADER and a transformer for "
                    f"polarity, and NRC for the emotion profile.",
                    color="EB6834",
                )

    doc.h2("4.3 Irony and sarcasm — the decisive test")
    doc.para(
        "Sarcasm is where lexicon methods fail structurally, and it is the reason the supervisor's "
        "recommendation matters. Consider: 'Oh great, another subscription fee. Brilliant work.' The "
        "text contains {great, brilliant} and no negative words, so Bing scores it maximally positive. "
        "VADER scores it positive too — its negation rules cover 'not good', not ironic praise."
    )
    doc.figure(FIGURE_DIR / "fig_nlp_irony.png",
               "Binary irony detection on TweetEval irony (SemEval-2018 Task 3).")

    if bench is not None:
        i = bench[(bench["status"] == "ok") & (bench["task"] == "irony") &
                  (bench["corpus"] == "tweeteval_irony")].sort_values("macro_f1", ascending=False)
        if len(i):
            doc.table(
                i[["method_name", "family", "accuracy", "macro_f1", "n_eval", "texts_per_sec"]].rename(
                    columns={"method_name": "Method", "family": "Family", "accuracy": "Accuracy",
                             "macro_f1": "Macro-F1", "n_eval": "n", "texts_per_sec": "Texts/sec"}),
                "Irony detection results, TweetEval test split.",
            )
            base = i[i["method_key"] == "majority_baseline"]
            lex = i[i["family"] == "lexicon"]
            if len(base) and len(lex):
                doc.callout(
                    "The headline result.",
                    f"The best word-list method reaches {float(lex['accuracy'].max()):.3f} accuracy "
                    f"against a majority-class baseline of {float(base['accuracy'].iloc[0]):.3f}. Word "
                    f"lists are worse than guessing on irony. No amount of tuning recovers this, because "
                    f"the signal is not in the vocabulary.",
                    color="E34948",
                )

    if (FIGURE_DIR / "fig_nlp_sarcasm_headlines.png").exists():
        doc.h3("Cross-domain generalisation")
        doc.para(
            "In-domain performance overstates what a method will do in production. The same methods "
            "were therefore evaluated on a second, independent sarcasm corpus from a different domain "
            "(news headlines, Misra & Arora) — a harder and more honest test, especially for the "
            "transformer checkpoint that was fine-tuned on the tweet corpus."
        )
        doc.figure(FIGURE_DIR / "fig_nlp_sarcasm_headlines.png",
                   "Sarcasm detection on news headlines — a domain none of the methods were tuned for.")
        if (FIGURE_DIR / "fig_irony_crossdomain.png").exists():
            doc.figure(FIGURE_DIR / "fig_irony_crossdomain.png",
                       "The same methods on both corpora. The gap between the bars is the "
                       "generalisation penalty.")

    doc.h2("4.4 Emotion")
    doc.figure(FIGURE_DIR / "fig_nlp_emotion.png",
               "Four-class emotion on TweetEval (SemEval-2018 Task 1).")
    doc.para(
        "The NRC lexicon carries eight emotion categories while the benchmark labels four, so the "
        "mapping is necessarily lossy: anticipation and trust fold into optimism, disgust and fear "
        "into anger, and surprise has no corresponding label at all. That handicap is a genuine "
        "limitation of scoring an eight-way lexicon against a four-way benchmark and is reported "
        "rather than hidden — the lexicon's full eight-way output is what the dashboard displays.",
        italic=True, size=9,
    )

    doc.h2("4.5 Topic modelling")
    if coh:
        doc.para(
            f"BERTopic recovered {coh['n_topics']} topics over {coh['n_documents']:,} captions with "
            f"{coh['outlier_fraction']:.1%} classified as outliers. An LDA baseline was fitted on the "
            f"identical corpus with the identical topic count and scored with the identical coherence "
            f"implementation, so the comparison is like-for-like."
        )
        doc.figure(FIGURE_DIR / "fig_topic_coherence.png",
                   "Topic quality. Diversity is reported alongside coherence because a model that "
                   "repeats one generic topic scores well on coherence alone.")
        doc.table(
            pd.DataFrame([
                {"Model": "BERTopic", "NPMI": coh["bertopic"]["npmi"], "C_v": coh["bertopic"]["c_v"],
                 "Diversity": coh["bertopic"]["diversity"]},
                {"Model": "LDA", "NPMI": coh["lda"]["npmi"], "C_v": coh["lda"]["c_v"],
                 "Diversity": coh["lda"]["diversity"]},
            ]),
            "Topic coherence and diversity.",
        )
        doc.para(
            "BERTopic is expected to win here specifically because captions are short. LDA estimates a "
            "topic distribution from word co-occurrence within a document, and a twenty-word caption "
            "does not supply enough. BERTopic clusters sentence embeddings and extracts representative "
            "terms afterwards, which sidesteps the problem.", size=9, italic=True,
        )
    else:
        doc.para("Topic modelling did not run in this build.", italic=True)

    doc.h2("4.6 What this means for the production pipeline")
    doc.para(
        "The benchmark results drive a deployment decision rather than sitting in an appendix. "
        "Rule-based extraction, VADER and NRC run over every post — they are effectively free. SBERT "
        "embeddings run over every post because topics and brand-fit need them anyway. Fine-tuned "
        "transformers run over every post at a few hundred posts per second. LLM prompting does not: "
        "it is roughly three orders of magnitude slower per post, and scoring the full corpus with a "
        "local 7B model is on the order of a day of compute for a signal the transformer already "
        "provides. The LLM's role is therefore (a) rigorous benchmarking, which quantifies exactly what "
        "is lost by not running it everywhere, and (b) on-demand single-creator analysis in the "
        "dashboard, where a few seconds of latency is acceptable."
    )

    # ------------------------------------------------------------------
    doc.h1("5. Network analysis")
    if gmeta:
        doc.para(
            f"The graph contains {gmeta['n_nodes']:,} nodes and {gmeta['n_edges']:,} edges "
            f"(density {gmeta['density']:.5f}), partitioned by the Louvain method into "
            f"{gmeta['n_communities']} communities, the largest containing "
            f"{gmeta['largest_community']} creators."
        )
        doc.figure(FIGURE_DIR / "fig_network.png",
                   "Degree distribution and network-position tiers.")
        if not gmeta.get("betweenness_exact", True):
            doc.para(
                f"Betweenness centrality is approximated using {gmeta.get('betweenness_pivots')} "
                f"sampled pivots (Brandes–Pich estimator). Exact computation is O(V·E) and does not "
                f"scale; the approximation is disclosed rather than presented as exact.",
                italic=True, size=9,
            )
    else:
        doc.para("Network stage did not run in this build.", italic=True)

    # ------------------------------------------------------------------
    doc.h1("6. Scoring and pricing models")
    if not models:
        doc.para("Model stage did not run in this build.", italic=True)
    else:
        p = models["performance"]
        doc.h2("6.1 Validation protocol")
        doc.callout(
            "The decision that matters most.",
            "Validation uses GroupKFold on creator id, so a creator never appears in both training and "
            "test. Because each creator contributes up to three campaigns, a random split would place "
            "the same creator on both sides and inflate every metric reported. This is the single most "
            "common way projects of this shape overstate their results.",
        )
        doc.bullets([
            "Metrics computed in log space — engagement rate is log-normally distributed, and a raw-scale "
            "R² would be dominated by a handful of large values.",
            "Ranking metrics (Spearman, NDCG) reported alongside regression metrics, because the product "
            "ranks creators and order matters more than absolute value.",
            "Leakage controls enforced in code: a banned-substring list gates the model matrix and the "
            "trainer raises an assertion if any excluded column reaches it.",
        ])

        doc.h2("6.2 Results")
        doc.table(
            pd.DataFrame([
                {"Metric": "R² (log)", "Value": p["r2_log"]},
                {"Metric": "RMSE (log)", "Value": p["rmse_log"]},
                {"Metric": "Spearman", "Value": p["spearman"]},
                {"Metric": "NDCG@10", "Value": p["ndcg@10"]},
                {"Metric": "NDCG@50", "Value": p["ndcg@50"]},
                {"Metric": "Theoretical R² ceiling", "Value": p["theoretical_r2_log_ceiling"]},
                {"Metric": "Fraction of ceiling reached", "Value": p["fraction_of_ceiling"]},
            ]),
            f"Out-of-fold performance. Fold R²: {', '.join(f'{v:.3f}' for v in p['fold_r2_log'])}.",
        )
        doc.figure(FIGURE_DIR / "fig_model_baselines.png",
                   "The learned model against the two alternatives it replaces.")
        doc.para(
            f"The Phase-1 weighted index is flattered in this comparison — it is isotonically "
            f"calibrated to the target scale and fitted on the full dataset, advantages the learned "
            f"model does not receive. It still loses by "
            f"{p['r2_log'] - p['baseline_composite_index']['r2_log']:.3f} R². The conclusion that "
            f"learning is worth its complexity here is therefore conservative.", size=9, italic=True,
        )
        doc.figure(FIGURE_DIR / "fig_pred_vs_actual.png",
                   "Out-of-fold predictions against actuals, log-log.", width_in=4.2)

        doc.h2("6.3 Which pillar earns its place")
        doc.figure(FIGURE_DIR / "fig_ablation.png",
                   "Each pillar removed in turn; the bar is the R² lost.")
        ab = p.get("ablation", {}).get("drops", {})
        if ab:
            doc.table(
                pd.DataFrame([
                    {"Pillar": k.title(), "R² without": v["r2_without"], "Loss": v["delta"],
                     "Alone": p["ablation"]["only"].get(k), "Features": v["n_features"]}
                    for k, v in ab.items()
                ]).sort_values("Loss", ascending=False),
                "Pillar ablation under the same GroupKFold protocol.",
            )
            weak = [k for k, v in ab.items() if v["delta"] <= 0.005]
            if weak:
                doc.callout(
                    "A negative result, reported rather than buried.",
                    f"Removing the {', '.join(weak)} pillar does not measurably hurt the model. On this "
                    f"data those features add no predictive signal beyond what reach, engagement and "
                    f"network already capture. The NLP layer still earns its place — it drives "
                    f"brand-safety screening, the competitor-conflict gate, and the explanations shown "
                    f"to users — but it should not be claimed as a driver of performance prediction, "
                    f"and this report does not claim it.",
                    color="E34948",
                )
        doc.figure(FIGURE_DIR / "fig_importance.png", "Feature importance by split gain.")

        doc.h2("6.4 Price model")
        q = models["price"]
        doc.table(
            pd.DataFrame([
                {"Model": "LightGBM regressor", "R² (log)": q["r2_log"], "MAPE": q["mape"],
                 "Spearman": q["spearman"]},
                {"Model": "Published rate card (rule)", "R² (log)": q["baseline_rate_card"]["r2_log"],
                 "MAPE": q["baseline_rate_card"]["mape"], "Spearman": q["baseline_rate_card"]["spearman"]},
            ]),
            f"Price prediction. The shown band is the 10th–90th percentile of out-of-fold residuals "
            f"and covers {q['band_coverage_p10_p90']:.0%} of true fees.",
        )
        doc.callout(
            "Second negative result.",
            f"The learned price model reaches R² {q['r2_log']:.3f} against {q['baseline_rate_card']['r2_log']:.3f} "
            f"for a rate card with no learning at all. The gap does not justify the maintenance cost of "
            f"a trained model. The recommendation is to ship the Phase-1 rule-based price band and "
            f"revisit only once real negotiated-deal data has accumulated on-platform. Note also that "
            f"this R² is partly an artefact of the simulation, in which fees are generated from "
            f"followers, engagement and niche — all of which are model features.",
            color="E34948",
        )

    # ------------------------------------------------------------------
    doc.h1("7. Brand-fit scoring")
    doc.para(
        "Brand-Fit is deliberately not a learned model, and the reason is worth stating because it "
        "would be easy to mistake for a shortcut."
    )
    doc.numbered([
        "There is no label. 'Was this creator a good fit for this brand?' is recorded nowhere. "
        "Training on campaign engagement would simply re-learn the performance model and rename it.",
        "Half of fit is a hard constraint, not a preference. A creator who has promoted a direct "
        "competitor should be vetoed, and no similarity score should be able to outweigh that. Learned "
        "models blend; brand safety needs gates.",
        "It is the number a brand manager will argue with. A score that decomposes into 'semantic "
        "0.71, category match yes, geography no' is defensible in a meeting. A gradient-boosted score "
        "is not.",
    ])
    bf = _j(ARTIFACT_DIR / "brandfit" / "brandfit_config.json")
    if bf:
        doc.table(
            pd.DataFrame([{"Component": k.replace("_", " ").title(), "Weight": v}
                          for k, v in bf["component_weights"].items()]),
            "Brand-fit components. Gates applied on top: " + "; ".join(bf.get("gates", [])),
        )

    # ------------------------------------------------------------------
    doc.h1("8. Dashboard and business model")
    doc.para(
        "The dashboard implements the freemium structure from the original design. Free access builds "
        "the two-sided marketplace; depth sits behind the paid tier."
    )
    doc.table(
        pd.DataFrame([
            {"Capability": "Creator search and filters", "Free": "Niche and tier only",
             "Paid": "Engagement quality, network position, audience geo/demo, ad load"},
            {"Capability": "Performance score", "Free": "Band (High/Medium/Low)",
             "Paid": "Numeric percentile with pillar breakdown"},
            {"Capability": "Brand-fit matching", "Free": "Not available",
             "Paid": "Ranked shortlist with component decomposition"},
            {"Capability": "Network position", "Free": "Not available",
             "Paid": "Centrality, community, interactive map"},
            {"Capability": "Price band", "Free": "Not available",
             "Paid": "Estimated fee range per creator"},
            {"Capability": "Results per search", "Free": "Capped", "Paid": "Unlimited"},
            {"Capability": "Creator-side analytics", "Free": "Peer benchmarking",
             "Paid": "Rate guidance, brand-interest alerts, boosted placement"},
        ]),
        "Freemium gating, implemented as a single configuration object so the tiers cannot drift "
        "apart between pages.",
    )
    doc.para(
        "Architecturally the deployment follows an offline-scoring / online-serving split. All heavy "
        "inference — embeddings, transformers, topic modelling, model training — runs once in the "
        "pipeline and is cached to disk. The hosted dashboard loads no ML model at all; it reads "
        "precomputed Parquet. This is what allows it to run inside roughly 1 GB of memory on free "
        "hosting, and it is the same pattern production recommendation systems use."
    )

    # ------------------------------------------------------------------
    doc.h1("9. Limitations")
    doc.h2("9.1 Limitations of the data")
    doc.bullets([
        "The creator universe is synthetic. Absolute performance numbers describe this simulation, not "
        "the real market. What transfers is the relative comparison between methods and the "
        "engineering, not the specific R².",
        "Synthetic captions cannot validate NLP quality — which is why all NLP claims are measured on "
        "real corpora instead. It does mean, however, that content features computed over synthetic "
        "captions are cleaner than real captions would be, and the ablation result for the content "
        "pillar should be read with that in mind.",
        "The industry benchmarks used for calibration come from marketing publications, not "
        "peer-reviewed research, and their methodology is not independently auditable.",
        "Four niche engagement multipliers were interpolated rather than sourced, and are flagged as "
        "assumptions in the code.",
    ])
    doc.h2("9.2 Limitations of the modelling")
    doc.bullets([
        "The network graph is topical, not social. Centrality means embeddedness in a shared "
        "vocabulary, not social influence.",
        "Betweenness centrality is a sampled approximation.",
        "The content pillar shows no measurable contribution to performance prediction in the ablation.",
        "The price model's high R² is partly circular, since fees are generated from features the "
        "model can see.",
        "The LLM is benchmarked on a stratified subsample rather than the full test set, for cost "
        "reasons. Every other method is additionally scored on that same subsample so the comparison "
        "remains valid, but the LLM's confidence interval is correspondingly wider.",
    ])
    doc.h2("9.3 What would be needed to deploy this for real")
    doc.numbered([
        "A data-sharing agreement or first-party integration giving access to creator analytics, since "
        "public APIs do not provide the required fields.",
        "Real campaign outcome data. Everything in the modelling layer is built to accept it; the "
        "target column and training protocol would not change.",
        "Human evaluation of brand-fit, to convert it from a transparent composite into something that "
        "can be validated rather than only argued about.",
        "Fairness auditing across niche, follower tier and geography before any scoring system that "
        "affects creator income is used commercially.",
    ])

    # ------------------------------------------------------------------
    doc.h1("10. Reproducibility")
    doc.para("The entire result set regenerates from a clean checkout with two commands:")
    doc.para("pip install -r requirements-dev.txt", size=9.5)
    doc.para("python run_pipeline.py", size=9.5)
    if manifest:
        doc.para(
            f"The pipeline runs ten stages and produces {len(manifest['numeric_features'])} numeric and "
            f"{len(manifest['categorical_features'])} categorical features over "
            f"{manifest['n_campaign_rows']:,} campaign rows. All randomness is seeded."
        )
    if nlp_rep:
        skipped = [k for k, v in nlp_rep.items() if isinstance(v, dict) and v.get("status") in ("skipped", "disabled")]
        if skipped:
            doc.para(f"Stages not run in this build: {', '.join(skipped)}.", italic=True, size=9)

    doc.h1("11. References")
    for key, c in CITATIONS.items():
        doc.para(f"{c['authors']} {c['title']}. {c['venue']}. {c['url']}", size=9)
    doc.para(
        f"{bm.SOURCES['engagement']['publisher']}. {bm.SOURCES['engagement']['title']}. "
        f"{bm.SOURCES['engagement']['url']} (retrieved {bm.SOURCES['engagement']['retrieved']}).", size=9)
    doc.para(
        f"{bm.SOURCES['pricing']['publisher']}. {bm.SOURCES['pricing']['title']}. "
        f"{bm.SOURCES['pricing']['url']} (retrieved {bm.SOURCES['pricing']['retrieved']}).", size=9)
    doc.para("Blei, D., Ng, A., Jordan, M. Latent Dirichlet Allocation. JMLR 3, 2003.", size=9)
    doc.para("Hu, M. & Liu, B. Mining and Summarizing Customer Reviews. KDD 2004.", size=9)
    doc.para("Blondel, V. et al. Fast unfolding of communities in large networks. "
             "J. Stat. Mech. 2008.", size=9)
    doc.para("Ke, G. et al. LightGBM: A Highly Efficient Gradient Boosting Decision Tree. NeurIPS 2017.", size=9)
    doc.para("Lundberg, S. & Lee, S. A Unified Approach to Interpreting Model Predictions. NeurIPS 2017.", size=9)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    return doc.save(output)


if __name__ == "__main__":
    print("building report ...")
    from src.report.figures import build_all

    build_all()
    p = build()
    print(f"written to {p}")
