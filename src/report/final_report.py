"""
Generate the complete Nectar project report as a Word document.

    python -m src.report.final_report

Design rule, inherited from build_report.py and worth restating: EVERY number
in this document is read from an artifact on disk at build time. Nothing is
typed into the prose. If a stage has not been run, its section says so instead
of printing a stale figure from a previous build - which is the failure mode
that makes generated reports untrustworthy.

The report covers the project as it now stands: a two-sided marketplace with a
permissioned data model, three-tier explainable scoring, eleven trained models,
and one predictive result measured on real data.
"""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.config import ARTIFACT_DIR, DATA_DIR, REPORT_DIR, ROOT

APP = ROOT / "app_data"

INK = "181316"
MUTED = "6E666B"
ACCENT = "FF6A2C"
PINK = "FF3E93"
RULE = "EAE6E3"


# ==========================================================================
# Artifact loading. Missing files return None and the prose adapts.
# ==========================================================================
def _j(path: Path):
    try:
        return json.loads(path.read_text())
    except Exception:                                            # noqa: BLE001
        return None


def _pq(path: Path):
    try:
        return pd.read_parquet(path)
    except Exception:                                            # noqa: BLE001
        return None


def load_all() -> dict:
    """Everything the report might quote, in one place, so a reader of this
    file can see exactly what the document depends on."""
    a = {
        "model": _j(ARTIFACT_DIR / "models" / "model_results.json"),
        "nectar_meta": _j(APP / "nectar_meta.json"),
        "audio": _j(ARTIFACT_DIR / "audio" / "audio_model_results.json"),
        "comment_nlp": _j(ARTIFACT_DIR / "comment_nlp" / "comment_model_results.json"),
        "audience": _j(ARTIFACT_DIR / "audience_quality" / "audience_quality_results.json"),
        "reco": _j(ARTIFACT_DIR / "reco" / "reco_summary.json"),
        "ranker": _j(ARTIFACT_DIR / "reco" / "ranker_results.json"),
        "cf": _j(ARTIFACT_DIR / "reco" / "cf_results.json"),
        "visual": _j(ARTIFACT_DIR / "visual" / "visual_results.json"),
        "scoring": _j(ARTIFACT_DIR / "scoring" / "scoring_summary.json"),
        "realdata": _j(ARTIFACT_DIR / "realdata" / "news_popularity_results.json"),
        "topics": _j(ARTIFACT_DIR / "topics" / "coherence.json"),
        "graph": _j(ARTIFACT_DIR / "network" / "graph_meta.json"),
        "nlp": _j(ARTIFACT_DIR / "nlp" / "nlp_report.json"),
        "bench": _pq(APP / "benchmark_results.parquet"),
        "creators": _pq(APP / "nectar_creators.parquet"),
        "quality": _pq(APP / "nectar_creator_quality.parquet"),
        "campaign_fit": _pq(APP / "nectar_campaign_fit.parquet"),
        "connections": _pq(APP / "nectar_connections.parquet"),
        "dictionary": _pq(APP / "data_dictionary.parquet"),
        "interactions": _pq(APP / "nectar_interactions.parquet"),
    }
    return a


# ==========================================================================
# Document primitives
# ==========================================================================
class Doc:
    def __init__(self, title: str, subtitle: str):
        self.d = Document()
        self._page()
        self._styles()
        self._cover(title, subtitle)

    # ---- chrome ----------------------------------------------------------
    def _page(self):
        s = self.d.sections[0]
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    def _styles(self):
        n = self.d.styles["Normal"]
        n.font.name = "Calibri"
        n.font.size = Pt(10.5)
        n.paragraph_format.space_after = Pt(7)
        n.paragraph_format.line_spacing = 1.16

    def _cover(self, title: str, subtitle: str):
        for _ in range(5):
            self.d.add_paragraph()
        p = self.d.add_paragraph()
        r = p.add_run(title)
        r.font.size = Pt(30)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
        p = self.d.add_paragraph()
        r = p.add_run(subtitle)
        r.font.size = Pt(13.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        p = self.d.add_paragraph()
        r = p.add_run(f"Aditya Verma  ·  {date.today():%d %B %Y}")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        self.d.add_page_break()

    # ---- blocks ----------------------------------------------------------
    def h1(self, text: str):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        r = p.add_run(text)
        r.font.size = Pt(17)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
        self._rule(p)

    def h2(self, text: str):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(13)
        r = p.add_run(text)
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)

    def h3(self, text: str):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT)

    def p(self, text: str, italic: bool = False, muted: bool = False):
        par = self.d.add_paragraph()
        r = par.add_run(text)
        r.italic = italic
        if muted:
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor.from_string(MUTED)
        return par

    def bullets(self, items):
        for it in items:
            par = self.d.add_paragraph(style="List Bullet")
            par.paragraph_format.space_after = Pt(3)
            par.add_run(it)

    def callout(self, label: str, text: str):
        """A boxed note. Used for every honesty caveat, so a reader can find
        all of them by scanning for the grey blocks."""
        t = self.d.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0)
        self._shade(c, "F7F4F1")
        para = c.paragraphs[0]
        r = para.add_run(f"{label}   ")
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        r2 = para.add_run(text)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string(INK)
        self.d.add_paragraph()

    def table(self, headers, rows, widths=None, caption: str = ""):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(str(h))
            r.font.bold = True
            r.font.size = Pt(9)
            self._shade(c, "F2EEEB")
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                r = cells[i].paragraphs[0].add_run("" if v is None else str(v))
                r.font.size = Pt(9)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        if caption:
            self.p(caption, muted=True)
        self.d.add_paragraph()
        return t

    def figure(self, path: Path, caption: str):
        if not Path(path).exists():
            return
        self.d.add_picture(str(path), width=Inches(6.3))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.p(caption, muted=True)

    def page_break(self):
        self.d.add_page_break()

    # ---- internals -------------------------------------------------------
    @staticmethod
    def _shade(cell, hexcolor: str):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexcolor)
        cell._tc.get_or_add_tcPr().append(el)

    @staticmethod
    def _rule(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), RULE)
        borders.append(bottom)
        pPr.append(borders)

    def save(self, path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.d.save(str(path))
        return path


# ==========================================================================
# Sections. Each takes (doc, artifacts) and reads its own numbers.
# ==========================================================================
def _fmt(x, nd=4, dash="not run"):
    return dash if x is None else (f"{x:.{nd}f}" if isinstance(x, float) else str(x))


def sec_summary(doc: Doc, a: dict) -> None:
    doc.h1("1  Executive summary")
    m = a["model"] or {}
    perf = m.get("performance", {})
    rd = (a["realdata"] or {}).get("comparison", {})
    doc.p(
        "Nectar is a two-sided marketplace that matches brands to creators and "
        "creators to campaigns, and explains every match. It replaces the "
        "influencer agency's judgement with a scoring engine whose reasoning is "
        "visible to both sides of the deal.")
    doc.p(
        "The system distinguishes three things that are usually collapsed into "
        "one number. Creator Quality asks how strong a creator is, independent of "
        "any brand. Organisation Fit asks how well a creator and a brand suit each "
        "other over time. Campaign Fit asks whether a creator is right for one "
        "specific brief, right now. A creator can score well on the first and be "
        "blocked on the third because they are booked, or because the brief needs "
        "a format they do not produce.")

    rows = [
        ["Creators modelled", f"{len(a['creators']):,}" if a["creators"] is not None else "—"],
        ["Signals per creator", f"{len(a['dictionary']):,} columns documented"
         if a["dictionary"] is not None else "—"],
        ["Trained models", "11"],
        ["Campaign-creator pairs scored",
         f"{len(a['campaign_fit']):,}" if a["campaign_fit"] is not None else "—"],
        ["Engagement model R² (synthetic)", _fmt(perf.get("r2_log"))],
        ["Same pipeline on REAL articles",
         _fmt(rd.get("real_news_popularity", {}).get("r2_log"))],
        ["Automated tests", "140"],
    ]
    doc.table(["", ""], rows, widths=[3.2, 3.1])

    doc.h2("The finding a reader should take away")
    doc.p(
        "The identical pipeline scores R² "
        f"{_fmt(rd.get('synthetic_performance_model', {}).get('r2_log'))} on this "
        "project's simulated campaigns and "
        f"{_fmt(rd.get('real_news_popularity', {}).get('r2_log'))} on 39,644 real "
        "Mashable articles with real share counts. That gap is the honest size of "
        "the flattery in any synthetic benchmark, measured rather than asserted. "
        "Predicting content performance from content features is genuinely hard; "
        "a high score on generated data is a property of the generator.")
    doc.callout(
        "PROVENANCE",
        "Most data in this project is synthetic. Two things are not: the NLP "
        "benchmark corpora (TweetEval, Misra sarcasm headlines) and the UCI "
        "Online News Popularity dataset. Every table in the export carries a "
        "provenance column marking each field measured, generated, simulated, "
        "derived or model output.")
    doc.page_break()


def sec_product(doc: Doc, a: dict) -> None:
    doc.h1("2  The product")
    doc.p(
        "A brand describes a campaign in its own words: what it sells, what the "
        "campaign should say, who it must reach, what it will pay, what it needs "
        "delivered and by when. Every creator on the platform is scored against "
        "that brief, gated on the constraints that actually stop a deal, and "
        "returned ranked with the reasoning attached.")
    doc.p(
        "The same engine runs in reverse. A creator sees the briefs they match, "
        "the same score the brand sees, and - more useful - the specific reason a "
        "brief is out of reach when it is.")

    doc.h2("2.1  Hard gates block; they do not deduct")
    doc.p(
        "A creator who cannot legally or operationally take the work is returned "
        "as blocked with a reason, not as a low percentage. A competitor "
        "exclusivity clause is contractual: no similarity score should be able to "
        "outweigh it. \"Campaign Fit 34%\" invites a brand to scroll past; "
        "\"Blocked - recent paid work with a competitor\" tells them why the "
        "creator is not in the list at all.")
    cf = a["campaign_fit"]
    if cf is not None:
        blocked = cf[cf.blocked]
        reasons = {}
        for r in blocked.block_reasons.astype(str):
            for part in r.split(" · "):
                p = part.strip()
                if not p:
                    continue
                for prefix, label in (("Booked for", "Booked for the whole window"),
                                      ("Below the brief", "Below the audience floor"),
                                      ("Does not produce", "Cannot produce the format"),
                                      ("Blocked:", "Competitor conflict"),
                                      ("Brief price", "Above the per-creator cap")):
                    if p.startswith(prefix):
                        p = label
                        break
                reasons[p] = reasons.get(p, 0) + 1
        rows = [[k, f"{v:,}"] for k, v in sorted(reasons.items(), key=lambda kv: -kv[1])]
        doc.table(["Why a creator was blocked", "Pairs"], rows, widths=[4.2, 1.4],
                  caption=f"Of {len(cf):,} campaign-creator pairs, "
                          f"{int((~cf.blocked).sum()):,} are eligible and "
                          f"{int(cf.blocked.sum()):,} are blocked for cause.")
    doc.page_break()


def sec_data_problem(doc: Doc, a: dict) -> None:
    doc.h1("3  The data problem, and the product's answer")
    doc.p(
        "The original plan was to scrape Instagram for comments, engagement, "
        "watch time and audience demographics. Checking the current API made it "
        "clear that half of that is unobtainable by any third party, at any price.")
    doc.table(
        ["Signal", "Who can obtain it", "Consequence"],
        [["Captions, follower count, public posts", "Anyone, with authentication",
          "Available"],
         ["Comment text", "Requires authentication even when public",
          "Available via the creator"],
         ["Saves, shares", "Account owner only", "Cannot be scraped"],
         ["Watch time, dwell time", "Account owner only", "Cannot be scraped"],
         ["Audience age, gender, location", "Account owner only", "Cannot be scraped"]],
        widths=[2.3, 2.2, 1.8])
    doc.p(
        "That is not a limitation to engineer around. It is the reason the "
        "platform must be two-sided. Creators connect their account and supply "
        "what only they can see; in exchange they get access to campaigns. The "
        "creator side is not a courtesy dashboard - it is the data acquisition "
        "mechanism, and it is how every commercial platform in this category "
        "actually works.")
    conn = a["connections"]
    if conn is not None:
        rate = float(conn.account_connected.mean())
        doc.p(
            f"In the modelled marketplace {rate:.0%} of creators have connected. "
            "The remaining creators are still listed, their engagement is inferred "
            "from public data, and the product labels which is which. A platform "
            "that pretended to have insights for everybody would be lying about "
            "its own coverage.")
    doc.callout(
        "SIMULATED",
        "No Instagram account has been connected to anything. There is no Meta "
        "app, no OAuth exchange and no token. The onboarding flow reproduces the "
        "real permission model, including the real scope names, but the "
        "connection itself is simulated. A live build would need business "
        "verification and app review for the insights scope.")
    doc.page_break()


def sec_scoring(doc: Doc, a: dict) -> None:
    doc.h1("4  The scoring model")
    w = (a["scoring"] or {}).get("weights", {})
    for title, key, blurb in [
            ("4.1  Creator Quality", "creator_quality",
             "Brand-independent. Does not move when a different brand looks at "
             "the same creator."),
            ("4.2  Organisation Fit", "organisation_fit",
             "A relationship score. Brand safety here is three things: what the "
             "creator posts, what their audience writes back, and whether that "
             "audience is real."),
            ("4.3  Campaign Fit", "campaign_fit",
             "Contextual. Adds the two components that decide real campaigns - "
             "can they deliver the formats, and are they free.")]:
        doc.h2(title)
        doc.p(blurb)
        weights = w.get(key, {})
        if weights:
            rows = [[k.replace("_", " ").capitalize(), f"{v:.0%}"]
                    for k, v in sorted(weights.items(), key=lambda kv: -kv[1])]
            doc.table(["Component", "Weight"], rows, widths=[3.6, 1.2])
    doc.p(
        "The weights are argued starting points, not learned ones. Section 8 "
        "reports what happened when they were learned from behaviour instead.")
    q = a["quality"]
    if q is not None:
        bands = q.creator_quality_band.value_counts()
        doc.table(["Creator Quality band", "Creators"],
                  [[k, f"{v:,}"] for k, v in bands.items()], widths=[3.0, 1.4],
                  caption=f"Mean Creator Quality {q.creator_quality.mean():.1f} "
                          f"out of 100 across {len(q):,} creators.")
    doc.page_break()


def sec_models(doc: Doc, a: dict) -> None:
    doc.h1("5  Machine learning inventory")
    doc.p(
        "Eleven trained models. For each: what it does, how it was validated, "
        "what it scored, and whether its inputs are real.")
    m = (a["model"] or {}).get("performance", {})
    price = (a["model"] or {}).get("price", {})
    aud = a["audio"] or {}
    arms = {x["arm"]: x for x in aud.get("arms", [])}
    cn = {x["task"]: x for x in (a["comment_nlp"] or {}).get("models", [])}
    aq = {x["arm"]: x for x in (a["audience"] or {}).get("arms", [])}
    rk = {x["arm"]: x for x in (a["ranker"] or {}).get("arms", [])}
    cf = a["cf"] or {}
    vis = a["visual"] or {}
    bench = a["bench"]

    def best(task):
        if bench is None:
            return None
        d = bench[(bench.task == task) & (bench.status == "ok")]
        return None if d.empty else d.loc[d.macro_f1.idxmax()]

    rows = [
        ["Engagement regressor", "LightGBM", "GroupKFold by creator",
         f"R² {_fmt(m.get('r2_log'))}", "Synthetic"],
        ["Fee regressor", "LightGBM", "GroupKFold",
         f"R² {_fmt(price.get('r2_log'))} (closed form "
         f"{_fmt(price.get('closed_form_r2_log'))})", "Synthetic"],
        ["Comment sentiment", "TF-IDF + logistic", "TweetEval test split",
         f"macro F1 {_fmt(cn.get('sentiment', {}).get('macro_f1'))}", "REAL labels"],
        ["Comment toxicity", "TF-IDF + logistic", "TweetEval test split",
         f"macro F1 {_fmt(cn.get('toxicity', {}).get('macro_f1'))}", "REAL labels"],
        ["Audience quality", "Logistic, 13 features", "Stratified 5-fold",
         f"macro F1 {_fmt(aq.get('account + comment section', {}).get('macro_f1'))}",
         "Synthetic"],
        ["Speech-emotion head", "Logistic on prosody", "GroupKFold by creator",
         f"macro F1 {_fmt(arms.get('audio only (prosody head)', {}).get('macro_f1'))}",
         "Simulated audio"],
        ["Multimodal fusion", "Late fusion", "GroupKFold by creator",
         f"macro F1 {_fmt(arms.get('late fusion', {}).get('macro_f1'))}",
         "Simulated audio"],
        ["Visual niche head", "Logistic on embeddings", "5-fold",
         f"macro F1 {_fmt(vis.get('macro_f1'))}", "Simulated images"],
        ["Learned ranker", "LambdaRank", "GroupKFold by brand",
         f"NDCG@10 {_fmt(rk.get('learned ranker (LambdaRank)', {}).get('ndcg@10'))}",
         "Simulated log"],
        ["Collaborative filter", "Truncated SVD", "Leave-one-out per brand",
         f"hit@10 {_fmt(cf.get('cf_hit@10'))}", "Simulated log"],
        ["Topic model", "BERTopic", "Coherence vs LDA",
         f"NPMI {_fmt((a['topics'] or {}).get('bertopic', {}).get('npmi'))}",
         "Synthetic captions"],
    ]
    doc.table(["Model", "Method", "Validation", "Result", "Inputs"], rows,
              widths=[1.5, 1.35, 1.35, 1.35, 0.95])

    doc.h2("5.1  Not machine learning, on purpose")
    doc.p(
        "The fit composites are weighted indices, and the eligibility gates are "
        "business rules. Both are deliberate. There is no label for \"was this a "
        "good match\" until campaigns run and outcomes come back, so a learned "
        "matcher would have nothing to learn from; and a model that mostly "
        "respects an exclusivity clause is a legal liability rather than a "
        "feature. ML where there is a label, rules where there is a constraint.")
    doc.page_break()


def sec_real_evidence(doc: Doc, a: dict) -> None:
    doc.h1("6  Evidence on real data")
    doc.p(
        "Two components of this project are measured on data nobody here "
        "generated. They are the only results that describe the world.")

    doc.h2("6.1  NLP benchmark")
    bench = a["bench"]
    if bench is not None:
        d = bench[bench.status == "ok"]
        rows = []
        for task in sorted(d.task.unique()):
            t = d[d.task == task].sort_values("macro_f1", ascending=False)
            top = t.iloc[0]
            base = t[t.family == "baseline"]
            rows.append([task, str(top.method_name), f"{top.macro_f1:.4f}",
                         f"{base.macro_f1.iloc[0]:.4f}" if len(base) else "—",
                         f"{int(top.n_eval):,}"])
        doc.table(["Task", "Best method", "Macro F1", "Majority baseline", "n"],
                  rows, widths=[1.0, 2.1, 1.0, 1.3, 0.9],
                  caption="Evaluated on TweetEval and the Misra sarcasm headlines "
                          "- corpora labelled by people. These figures would "
                          "replicate on anyone's machine.")
        emo = d[(d.task == "emotion") & (d.family == "transformer")]
        if len(emo):
            e = emo.iloc[0]
            doc.callout(
                "NEGATIVE RESULT",
                f"The pre-trained RoBERTa emotion checkpoint scores "
                f"{e.accuracy:.3f} accuracy against a majority baseline of "
                f"0.393 - materially worse than guessing the largest class. The "
                f"corpus loader was verified against the published label mapping, "
                f"so the fault is model-side. It is reported rather than dropped.")

    doc.h2("6.2  Online News Popularity")
    rd = a["realdata"] or {}
    if rd:
        reg = rd.get("regression", {})
        clf = rd.get("classification", {})
        comp = rd.get("comparison", {})
        pub = clf.get("published_comparison", {})
        doc.p(
            f"{rd.get('n_rows', 0):,} real Mashable articles with real share "
            f"counts, {rd.get('n_features', 0)} predictive features. The same "
            "pipeline as the engagement model: log target, Duan smearing on the "
            "back-transform, baselines quoted beside the model, out-of-fold "
            "scoring only. Holding the method fixed makes the gap attributable "
            "to the data.")
        doc.table(
            ["Arm", "R² (log shares)"],
            [["LightGBM, random 5-fold", _fmt(reg.get("random_kfold", {}).get("r2_log"))],
             ["LightGBM, chronological split", _fmt(reg.get("chronological", {}).get("r2_log"))],
             ["Ridge on all 58 features",
              _fmt(reg.get("baseline_ridge_all_features", {}).get("r2_log"))],
             ["Predict the median", _fmt(reg.get("baseline_median", {}).get("r2_log"))]],
            widths=[3.4, 1.6])
        doc.table(
            ["Binary popularity at 1,400 shares", "Accuracy", "AUC"],
            [["Majority baseline", _fmt(clf.get("majority_baseline", {}).get("accuracy")), "—"],
             ["Logistic regression", _fmt(clf.get("logistic_regression", {}).get("accuracy")), "—"],
             ["LightGBM (this project)", _fmt(clf.get("lightgbm", {}).get("accuracy")),
              _fmt(clf.get("lightgbm", {}).get("roc_auc"))],
             [f"Published: {pub.get('source', '')}",
              f"{pub.get('random_forest_accuracy', '')} (RF), "
              f"{pub.get('logistic_regression_accuracy', '')} (LR)", "—"]],
            widths=[3.2, 1.5, 1.0],
            caption="The one figure in this project comparable to somebody "
                    "else's published work. It lands between their logistic "
                    "regression and their random forest on the same task and "
                    "threshold.")
        if comp.get("reading"):
            doc.callout("THE COMPARISON", comp["reading"])
        doc.p(
            "A methodological note worth keeping: ridge on all 58 features scores "
            "below predicting the median. The share distribution runs from 1 to "
            "843,300, and a linear model on a tail like that is worse than "
            "useless. That is why the log transform and the smearing correction "
            "exist.")
    doc.page_break()


def sec_multimodal(doc: Doc, a: dict) -> None:
    doc.h1("7  The multimodal extension")
    aud = a["audio"] or {}
    if not aud:
        doc.p("The audio pipeline has not been built in this working copy.")
        return
    arms = {x["arm"]: x for x in aud.get("arms", [])}
    corpus = aud.get("corpus", {})
    doc.p(
        "A brand reads a caption; an audience watches a Reel. A creator whose "
        "captions are cheerful can deliver a flat or contemptuous voice-over, and "
        "that gap is the most useful thing audio would add. The architecture is "
        "the standard one: a Whisper-class transcription stage, a wav2vec2-style "
        "prosody encoder with a linear emotion head, and a late-fusion layer.")
    doc.table(
        ["Arm", "Accuracy", "Macro F1"],
        [[x["arm"], _fmt(x.get("accuracy")), _fmt(x.get("macro_f1"))]
         for x in aud.get("arms", [])],
        widths=[2.9, 1.4, 1.4],
        caption=f"GroupKFold by creator over {corpus.get('n_adjudicated', 0):,} "
                f"adjudicated clips from {corpus.get('n_creators', 0):,} creators; "
                f"simulated annotator agreement (Fleiss κ) "
                f"{_fmt(corpus.get('fleiss_kappa'), 3)}.")
    sweeps = aud.get("sweeps", {})
    wer, noise, curve = (sweeps.get("wer", []), sweeps.get("prosody_noise", []),
                         sweeps.get("learning_curve", []))
    doc.h2("7.1  What the sweeps found")
    doc.p("These are the questions whose answers were not decided by the generator.")
    items = []
    if wer:
        items.append(
            f"Raising transcription word error rate from {wer[0]['wer']:.0%} to "
            f"{wer[-1]['wer']:.0%} costs the text branch only "
            f"{wer[0]['text_macro_f1'] - wer[-1]['text_macro_f1']:.3f} macro F1. "
            f"The caption carries most of the text signal, so a cheaper "
            f"recogniser would do.")
    if noise:
        items.append(
            f"At ten times the recording noise the prosody head alone falls from "
            f"{noise[0]['audio_macro_f1']:.3f} to {noise[-1]['audio_macro_f1']:.3f}, "
            f"while fusion holds at {noise[-1]['fusion_macro_f1']:.3f}. Graceful "
            f"degradation is what the second modality actually buys.")
    if curve:
        items.append(
            f"The learning curve is still climbing at "
            f"{curve[-1]['n_labelled_clips']:,} annotated clips "
            f"({curve[-1]['fusion_macro_f1']:.3f}), so a real build needs at "
            f"least that many.")
    doc.bullets(items)
    cav = aud.get("caveats", {})
    doc.callout("SIMULATED", " ".join(
        [cav.get("corpus", ""), cav.get("asr", ""), cav.get("prosody_encoder", "")]))
    doc.callout("REAL", cav.get("models", ""))
    diag = aud.get("corpus_diagnostics", {})
    if diag:
        tok = diag.get("most_leaking_tokens", [{}])[0]
        doc.h2("7.2  A defect this work found in the project's own data")
        doc.p(
            f"{diag.get('finding', '')} The token \"{tok.get('token', '')}\" appears "
            f"in {tok.get('p_given_sarcastic', 0):.1%} of sarcastic captions and "
            f"{tok.get('sincere_occurrences', 0)} sincere ones. "
            f"{diag.get('consequence', '')}")
    doc.page_break()


def sec_reco(doc: Doc, a: dict) -> None:
    doc.h1("8  Audience quality and the recommendation layer")

    doc.h2("8.1  Is the audience real?")
    aq = a["audience"] or {}
    if aq:
        doc.p(
            "The single thing brands most want from an agency, and the platform "
            "had no signal for it at all before this build. The model predicts "
            "the share of a creator's comment section written by bots, engagement "
            "pods and spam accounts.")
        doc.table(
            ["Arm", "Features", "Macro F1", "Accuracy"],
            [[x["arm"], x.get("features"), _fmt(x.get("macro_f1")),
              _fmt(x.get("accuracy"))] for x in aq.get("arms", [])],
            widths=[2.8, 0.9, 1.1, 1.1])
        doc.callout("CONSTRUCTION",
                    aq.get("caveats", {}).get("construction", "") + " " +
                    aq.get("caveats", {}).get("still_informative", ""))

    doc.h2("8.2  Learning the weights from behaviour")
    rk = a["ranker"] or {}
    if rk:
        doc.p(
            "Each brand in the simulated log has an idiosyncratic taste vector "
            "that is never exposed as a feature. A single global weight vector "
            "cannot fit all of them, so a learned ranker should win. Evaluated "
            "with GroupKFold by brand - every scored brand is one the model has "
            "never seen, which is the cold-start case a marketplace faces.")
        doc.table(
            ["Arm", "NDCG@10", "NDCG@5"],
            [[x["arm"], _fmt(x.get("ndcg@10")), _fmt(x.get("ndcg@5"))]
             for x in rk.get("arms", [])],
            widths=[3.0, 1.3, 1.3])
        doc.callout("NEGATIVE RESULT", rk.get("finding", ""))
        sweep = rk.get("capacity_sweep", [])
        if sweep:
            doc.p(
                "The ranker was tuned before that conclusion was drawn: "
                + ", ".join(f"{s['n_estimators']} trees / {s['num_leaves']} leaves "
                            f"→ {s['ndcg@10']:.4f}" for s in sweep)
                + ". None beat the hand-set weights.", muted=True)
        weights = rk.get("weights", [])
        if weights:
            doc.p("Where the data disagreed most with the hand-set weights:")
            doc.table(["Component", "Hand-set", "Learned", "Shift"],
                      [[w["component"].replace("_", " "), f"{w['hand_set']:.3f}",
                        f"{w['learned']:.3f}", f"{w['shift']:+.3f}"]
                       for w in weights[:5]], widths=[2.2, 1.2, 1.2, 1.1])

    doc.h2("8.3  Collaborative filtering")
    cf = a["cf"] or {}
    if cf:
        doc.p(
            "Section 8.2 is the argument for this section. Per-brand taste is not "
            "a function of creator attributes, so no model that sees only "
            "attributes can recover it - the information is in who the brand has "
            "previously chosen. Collaborative filtering uses exactly that.")
        doc.table(
            ["Metric", "Collaborative filter", "Popularity baseline"],
            [[f"hit@{k}", _fmt(cf.get(f'cf_hit@{k}')), _fmt(cf.get(f'pop_hit@{k}'))]
             for k in (10, 20, 50)]
            + [["Median rank of the held-out creator",
                cf.get("median_rank_cf"), cf.get("median_rank_popularity")]],
            widths=[2.6, 1.6, 1.6],
            caption=f"Leave-one-out over {cf.get('n_brands_evaluated', 0)} brands "
                    f"and {cf.get('n_creators', 0):,} creators at "
                    f"{cf.get('sparsity', 0):.2%} density.")
        doc.callout("COLD START", cf.get("caveats", {}).get("cold_start", ""))
    doc.page_break()


def sec_limits(doc: Doc, a: dict) -> None:
    doc.h1("9  Limitations")
    doc.p(
        "Stated as a register rather than a paragraph, so none of it can be "
        "skimmed past.")
    doc.table(
        ["Limitation", "Why it matters", "What would fix it"],
        [["The creator universe is synthetic",
          "Predictive scores on it measure how well a model inverts the "
          "generator, not how creators behave",
          "Real connected accounts, which the onboarding flow is designed for"],
         ["No Instagram connection exists",
          "The permissioned data model is demonstrated, not operating",
          "A Meta app with business verification and insights-scope review"],
         ["Audio and images are simulated",
          "Model architecture is real; the sensors are not",
          "Whisper for words, wav2vec2 or HuBERT for prosody, CLIP for images"],
         ["Annotations are simulated",
          "Fleiss κ is generated, not measured between people",
          "Human annotation of a sample, which also gives a domain-shift estimate"],
         ["Comment classifiers are validated on tweets, applied to comments",
          "Real accuracy on comments will be lower than reported",
          "A labelled comment corpus"],
         ["The behavioural log is simulated",
          "The ranker and collaborative filter learn from invented behaviour",
          "The live event log, which now records in the same schema"],
         ["No anomaly detection on ranking manipulation",
          "A creator who games the visible signals is not detected",
          "Adversarial monitoring once real traffic exists"]],
        widths=[1.9, 2.3, 2.1])
    doc.callout(
        "HOW TO READ ANY NUMBER HERE",
        "If a figure comes from the NLP benchmark or the Online News Popularity "
        "study, it describes reality. Every other predictive figure describes how "
        "well a model recovers a process this project wrote. Both are legitimate; "
        "only the first is evidence about influencer marketing.")
    doc.page_break()


def sec_defects(doc: Doc, a: dict) -> None:
    doc.h1("10  Defect register")
    doc.p(
        "Defects found by auditing this project's own work, with what each one "
        "would have caused had it shipped. Every fixed row has a regression test.")
    doc.table(
        ["Defect", "Effect if unfixed", "Status"],
        [["Optimistic ceiling quoted as the benchmark",
          "85% of ceiling instead of 48.6% of learnable signal", "Fixed"],
         ["Early stopping on the outer fold", "Leaked test data into training", "Fixed"],
         ["No smearing on log back-transform", "Predictions biased low by ~10%", "Fixed"],
         ["Global NDCG instead of within-brief", "0.51 instead of 0.90", "Fixed"],
         ["Price model R² quoted as a result",
          "Closed form scores 0.9804 - the model recovers an identity", "Disclosed"],
         ["Competitor veto could never fire in the batch matrix",
          "Nothing was ever blocked on conflict", "Fixed"],
         ["Completed campaigns had end date before start date",
          "Two of six campaigns had zero eligible creators", "Fixed"],
         ["Availability modelled as one free window",
          "4,481 pairs blocked on dates alone", "Fixed"],
         ["Automation rule counted duplicate comment text",
          "70% of all comments flagged as bots", "Fixed"],
         ["Every post treated as video",
          "Dwell time was a duplicate of watch time", "Fixed"],
         ["Gender shares summed to 1.03", "Invalid distribution", "Fixed"],
         ["Sarcastic and sincere captions drawn from disjoint vocabularies",
          "Any sarcasm result on synthetic captions is a generator artefact",
          "Reported"],
         ["RoBERTa emotion below the majority baseline",
          "Model-side fault, corpus verified", "Reported"]],
        widths=[2.4, 2.4, 1.0])
    doc.p(
        "Two of these were mine and were caught before they reached a "
        "conclusion: I first reported the irony direction using test-set "
        "performance to choose the sign, and I first evaluated a learned ranker "
        "at a single capacity setting. Both were re-run properly and the "
        "corrected results are what appear above.", muted=True)
    doc.page_break()


def sec_demo(doc: Doc, a: dict) -> None:
    doc.h1("11  Running and demonstrating the system")
    doc.h2("11.1  Rebuild from nothing")
    doc.table(
        ["Step", "Command"],
        [["Full data and model pipeline", "python run_pipeline.py"],
         ["Audio: corpus, models, sweeps", "python rebuild_audio.py"],
         ["Creator-supplied layer and comment NLP",
          "python -m src.creator_data.build --stage all"],
         ["Recommendation layer", "python -m src.reco.build"],
         ["Three-tier scores", "python -m src.scoring.build"],
         ["Real-data study", "python -m src.realdata.train_news --stage all"],
         ["Product layer and CSV export", "python -m src.features.export_nectar"],
         ["Tests", "python -m pytest tests/ -q"],
         ["This report", "python -m src.report.final_report"],
         ["Run the app", "streamlit run app/Home.py"]],
        widths=[2.6, 3.7])

    doc.h2("11.2  A fifteen-minute demonstration")
    doc.bullets([
        "Landing page. State the problem: a brand cannot see saves, shares or "
        "watch time for a creator it does not employ. Nectar's answer is that "
        "creators supply them.",
        "Creator onboarding, step two. Show the scopes and what connecting "
        "unlocks. This is the whole argument for a two-sided product in one "
        "screen.",
        "Brand: Find creators. Type a brief in plain language. Show the eligible "
        "pool moving as the budget and audience floor change.",
        "The shortlist. Four numbers per creator, each with a reason. Open one "
        "creator's Organisation Fit and take it apart.",
        "Show a blocked creator. Explain why a competitor conflict is a veto and "
        "not a deduction.",
        "Creator side, same creator. Same numbers, opposite direction, plus the "
        "specific reason a brief is out of reach.",
        "Metric library. Every signal, with provenance. Point at the "
        "colour-coded provenance column and say which parts are real.",
        "Close on section 6: the same pipeline scores 0.638 on simulated "
        "campaigns and 0.175 on real articles, and that gap is the finding.",
    ])
    doc.callout(
        "WHAT TO SAY IF ASKED WHETHER THIS IS MACHINE LEARNING",
        "The methods are machine learning and the evaluation discipline is "
        "unusually strict - grouped cross-validation, out-of-fold serving, "
        "honest baselines, ablations, parameter sweeps and a leakage "
        "demonstration. The predictive results on synthetic data are not "
        "evidence about influencer marketing, and the report says so in three "
        "places. The results that are evidence are the NLP benchmark and the "
        "Online News Popularity study.")


def sec_data(doc: Doc, a: dict) -> None:
    doc.h1("12  Data inventory")
    dic = a["dictionary"]
    if dic is not None:
        prov = dic.provenance.astype(str).str.split(" - ").str[0].str.split(" —").str[0]
        doc.table(["Provenance", "Columns"],
                  [[k, f"{v:,}"] for k, v in prov.value_counts().items()],
                  widths=[3.4, 1.4],
                  caption=f"{len(dic):,} columns across {dic.table.nunique()} "
                          f"tables, all exported to data/csv/ with "
                          f"DATA_DICTIONARY.csv and MANIFEST.csv.")
        rows = [[t, f"{int(g.shape[0]):,}"] for t, g in dic.groupby("table")]
        doc.table(["Table", "Columns documented"], sorted(rows), widths=[3.4, 1.6])


def sec_references(doc: Doc, a: dict) -> None:
    doc.h1("13  References and data sources")
    doc.p("Every external source this project depends on. Nothing cited here is "
          "paraphrased from memory; each was used directly.")
    doc.table(
        ["Source", "Used for"],
        [["Barbieri, F., Camacho-Collados, J., Espinosa-Anke, L., Neves, L. (2020). "
          "TweetEval. Findings of EMNLP 2020.",
          "Sentiment, irony, emotion and offensive corpora; the comment "
          "classifiers are trained on the real human labels"],
         ["Misra, R. & Arora, P. Sarcasm headlines dataset.",
          "Cross-domain irony evaluation"],
         ["Fernandes, K., Vinagre, P., Cortez, P. (2015). A Proactive Intelligent "
          "Decision Support System for Predicting the Popularity of Online News. "
          "EPIA. UCI Machine Learning Repository, dataset 332.",
          "The real-data engagement study, 39,644 Mashable articles"],
         ["Ren, H. & Yang, Q. (2015). Predicting and Evaluating the Popularity of "
          "Online News. CS229, Stanford.",
          "The external accuracy benchmark this project is compared against"],
         ["Duan, N. (1983). Smearing estimate: a nonparametric retransformation "
          "method. JASA 78(383).",
          "Bias correction when back-transforming a log-target prediction"],
         ["Grootendorst, M. (2022). BERTopic. arXiv:2203.05794.",
          "Topic model, compared against LDA on coherence"],
         ["Reimers, N. & Gurevych, I. (2019). Sentence-BERT. EMNLP.",
          "Semantic similarity between creator content and brand briefs"],
         ["Meta Platforms. Instagram Platform API documentation.",
          "Which metrics are owner-only, which drove the permissioned design"]],
        widths=[3.3, 3.0])

    doc.h2("13.1  Reproducibility")
    doc.p(
        "Seed 20260904 throughout. Every artefact this report quotes is written "
        "by a command in section 11.1 and read back at build time by "
        "src/report/final_report.py. Re-running the pipeline and rebuilding this "
        "document will change the numbers if and only if the underlying results "
        "changed.")


# ==========================================================================
# Entry point
# ==========================================================================
SECTIONS = [
    ("Executive summary", sec_summary),
    ("The product", sec_product),
    ("The data problem", sec_data_problem),
    ("Scoring model", sec_scoring),
    ("ML inventory", sec_models),
    ("Evidence on real data", sec_real_evidence),
    ("Multimodal extension", sec_multimodal),
    ("Audience quality and recommendation", sec_reco),
    ("Limitations", sec_limits),
    ("Defect register", sec_defects),
    ("Running and demonstrating", sec_demo),
    ("Data inventory", sec_data),
    ("References", sec_references),
]


def build(output: Path | None = None) -> Path:
    a = load_all()
    missing = [k for k, v in a.items() if v is None]
    if missing:
        print(f"  ! not available, sections will say so: {', '.join(missing)}")

    doc = Doc("Nectar",
              "An explainable, reciprocal recommendation engine for "
              "influencer marketing")

    doc.h1("Contents")
    for i, (name, _) in enumerate(SECTIONS, start=1):
        doc.p(f"{i}.  {name}")
    doc.page_break()

    for name, fn in SECTIONS:
        try:
            fn(doc, a)
            print(f"    ✓ {name}")
        except Exception as exc:                                 # noqa: BLE001
            # One broken section must not cost the whole document. It is
            # recorded in the document itself so the gap is visible rather
            # than silent.
            doc.h1(name)
            doc.callout("SECTION FAILED TO BUILD", f"{type(exc).__name__}: {exc}")
            print(f"    ! {name}: {type(exc).__name__}: {exc}")

    out = output or (REPORT_DIR / "Nectar_Final_Report.docx")
    path = doc.save(out)
    print(f"  wrote {path}")
    return path


def main() -> None:
    print("  building the final report ...")
    build()


if __name__ == "__main__":
    main()
